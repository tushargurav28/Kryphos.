"""
================================================================================
                    SECURITY RECONNAISSANCE REPORT GENERATOR
                           Recon Dashboard v4.1
================================================================================
"""

# =============================================================================
# SECTION 1: IMPORTS AND DEPENDENCIES
# =============================================================================

import json
import io
import os
import re
import sys
import csv
import logging
import hashlib
import xml.etree.ElementTree as ET
from datetime import datetime
from typing import List, Dict, Any, Optional
from collections import defaultdict
from pathlib import Path
from enum import Enum
import html as html_module

REPORTLAB_AVAILABLE = False
OPENPYXL_AVAILABLE = False
DOCX_AVAILABLE = False

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm, mm
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
    from reportlab.lib.colors import HexColor, Color as RLColor, black, white
    from reportlab.platypus import (
        SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer,
        PageBreak, Image, Flowable, KeepTogether,
        NextPageTemplate, PageTemplate, Frame, BaseDocTemplate
    )
    from reportlab.graphics.shapes import Drawing, Rect, String, Group
    from reportlab.graphics.charts.piecharts import Pie
    try:
        from reportlab.lib.validators import Auto
    except ImportError:
        Auto = None
    REPORTLAB_AVAILABLE = True
    print("ReportLab imported OK")
except ImportError as e:
    REPORTLAB_AVAILABLE = False
    print(f"ReportLab not available: {e}")

    class Drawing: pass
    class RLColor: pass
    class HexColor:
        def __init__(self, color): self.color = color
        def __str__(self): return self.color
    class ParagraphStyle: pass
    class TableStyle: pass
    class Table: pass
    class Paragraph: pass
    class Spacer: pass
    class PageBreak: pass
    class SimpleDocTemplate: pass
    Auto = None

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font as XLFont, Alignment as XLAlignment, PatternFill
    from openpyxl.utils import get_column_letter
    OPENPYXL_AVAILABLE = True
    print("OpenPyXL imported OK")
except ImportError as e:
    OPENPYXL_AVAILABLE = False
    print(f"OpenPyXL not available: {e}")

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
    print("Python-Docx imported OK")
except ImportError as e:
    DOCX_AVAILABLE = False
    print(f"Python-Docx not available: {e}")


# =============================================================================
# SECTION 2: CONFIGURATION CONSTANTS
# =============================================================================

class SeverityLevel(Enum):
    CRITICAL = "critical"
    HIGH = "high"
    MEDIUM = "medium"
    LOW = "low"
    INFO = "info"
    UNKNOWN = "unknown"


class ReportConfig:
    """All layout and style configuration in one place."""

    # --- Page (Letter = 612 x 792 pt) ---
    PAGE_SIZE    = letter if REPORTLAB_AVAILABLE else (612, 792)
    PAGE_WIDTH   = PAGE_SIZE[0]
    PAGE_HEIGHT  = PAGE_SIZE[1]
    MARGIN_TOP   = 0.45 * inch if REPORTLAB_AVAILABLE else 32
    MARGIN_BOTTOM= 0.45 * inch if REPORTLAB_AVAILABLE else 32
    MARGIN_LEFT  = 0.50 * inch if REPORTLAB_AVAILABLE else 36
    MARGIN_RIGHT = 0.50 * inch if REPORTLAB_AVAILABLE else 36
    # Usable = 612 - 36 - 36 = 540 pt = 7.5"
    USABLE_WIDTH = (PAGE_WIDTH - MARGIN_LEFT - MARGIN_RIGHT
                    if REPORTLAB_AVAILABLE else 540)

    # --- Fonts ---
    FONT_TITLE  = 20
    FONT_H1     = 13
    FONT_H2     = 11
    FONT_H3     = 10
    FONT_NORMAL = 9
    FONT_SMALL  = 8
    FONT_TINY   = 7
    TBL_HDR_SZ  = 9
    TBL_CELL_SZ = 8
    TBL_PAD     = 4

    # --- Colors ---
    C_PRIMARY    = HexColor('#1a252f')
    C_SECONDARY  = HexColor('#2c3e50')
    C_ACCENT     = HexColor('#2980b9')
    C_SUCCESS    = HexColor('#1e8449')
    C_WARNING    = HexColor('#d4ac0d')
    C_DANGER     = HexColor('#cb4335')
    C_INFO       = HexColor('#117a65')
    C_LIGHT      = HexColor('#f2f3f4')
    C_WHITE      = HexColor('#ffffff')
    C_GRAY       = HexColor('#85929e')
    C_GRAY_LT    = HexColor('#d5d8dc')
    C_DARK       = HexColor('#1a252f')
    C_CRITICAL   = HexColor('#922b21')
    C_CRIT_BG    = HexColor('#fadbd8')
    C_HIGH       = HexColor('#a04000')
    C_HIGH_BG    = HexColor('#fdebd0')
    C_MEDIUM     = HexColor('#9a7d0a')
    C_MED_BG     = HexColor('#fef9e7')
    C_LOW        = HexColor('#1a5276')
    C_LOW_BG     = HexColor('#eaf2ff')
    C_GRAYINFO   = HexColor('#616a6b')
    C_ROW_ALT    = HexColor('#f8f9fa')

    # --- Spacing (tight) ---
    SP_XS = 0.06 * inch if REPORTLAB_AVAILABLE else 4
    SP_SM = 0.10 * inch if REPORTLAB_AVAILABLE else 7
    SP_MD = 0.15 * inch if REPORTLAB_AVAILABLE else 11
    SP_LG = 0.22 * inch if REPORTLAB_AVAILABLE else 16

    # --- Limits ---
    MAX_HOSTS     = 25
    MAX_PORTS     = 35
    MAX_CH_CARDS  = 20

    DEFAULT_EXPORT_OPTIONS = {'include_recommendations': True}


# =============================================================================
# SECTION 3: UTILITIES
# =============================================================================

class LogUtils:
    _logger = None
    _initialized = False

    class _W:
        def __init__(self, lg): self._l = lg
        def info(self, m, **k): self._l.info(m, **k)
        def warning(self, m, **k): self._l.warning(m, **k)
        def error(self, m, **k): self._l.error(m, **k)
        def debug(self, m, **k): self._l.debug(m, **k)
        def exception(self, m, **k): self._l.exception(m, **k)
        def log_info(self, m, **k): self._l.info(m, **k)
        def log_warning(self, m, **k): self._l.warning(m, **k)
        def log_error(self, m, **k): self._l.error(m, **k)
        def log_debug(self, m, **k): self._l.debug(m, **k)
        def log_exception(self, m, exc=None):
            if exc: self._l.exception(f"{m}: {exc}")
            else:   self._l.exception(m)

    @classmethod
    def get_logger(cls, name='ReportGenerator'):
        if not cls._initialized: cls._initialize()
        if cls._logger is None:
            cls._logger = cls._W(logging.getLogger(name))
        return cls._logger

    @classmethod
    def _initialize(cls):
        if cls._initialized: return
        lg = logging.getLogger('ReportGenerator')
        lg.setLevel(logging.INFO)
        h = logging.StreamHandler(sys.stdout)
        h.setFormatter(logging.Formatter('%(asctime)s %(levelname)s %(message)s','%H:%M:%S'))
        if not lg.handlers: lg.addHandler(h)
        cls._initialized = True

    @classmethod
    def log_info(cls, m, **k): cls.get_logger().info(m, **k)
    @classmethod
    def log_error(cls, m, **k): cls.get_logger().error(m, **k)
    @classmethod
    def log_exception(cls, m, exc=None):
        l = cls.get_logger()
        if exc: l.exception(f"{m}: {exc}")
        else:   l.exception(m)


class TextUtils:
    _BOX = set('■●▲▼◆◇★☆○□▢▣▤▥▦▧▨▩▪▫▬▭▮▯▰▱▒▓─━│┃┄┅┆┇┈┉┊┋┌┐└┘├┤┬┴┼'
               '═║╒╓╔╕╖╗╘╙╚╛╜╝╞╟╠╡╢╣╤╥╦╧╨╩╪╫╬†‡•‣…‰′″‴‵‶‷‸‹›※‽‾')

    @staticmethod
    def safe_text(text, max_length=300, truncate=False):
        if text is None: return 'N/A'
        text = str(text)
        if not text.strip(): return 'N/A'
        text = ''.join(c for c in text if c not in TextUtils._BOX)
        if truncate and len(text) > max_length:
            return text[:max_length-3]+'...'
        return text.strip()

    @staticmethod
    def escape_html(text):
        return html_module.escape(str(text)) if text else ''

    @staticmethod
    def escape_pdf(text):
        if not text: return ''
        for c,e in [('&','&amp;'),('<','&lt;'),('>','&gt;'),('"','&quot;'),("'",'&apos;')]:
            text = text.replace(c,e)
        return text

    @staticmethod
    def get_status_text(status):
        m = {'configured':'Configured','missing':'Missing','not_found':'Not Found',
             'not_enabled':'Not Enabled','not_configured':'Not Configured',
             'complete':'Complete','verified':'Verified','review':'Review Required',
             'none':'None','minor':'Minor','awareness':'Awareness',
             'action':'Action Required','urgent':'Urgent'}
        return m.get(status.lower(), status.title())

    @staticmethod
    def sanitize_filename(fn):
        if not fn: return 'report'
        for c in '<>:"/\\|?*': fn=fn.replace(c,'_')
        return fn[:255].strip()


class DateUtils:
    @staticmethod
    def now(): return datetime.now()
    @staticmethod
    def get_timestamp(): return datetime.now().strftime('%Y%m%d_%H%M%S')


# =============================================================================
# SECTION 4: STYLES
# =============================================================================

class ReportStyles:
    def __init__(self):
        if REPORTLAB_AVAILABLE:
            self.styles = getSampleStyleSheet()
            self._create()
        else:
            self.styles = {}

    def _add(self, name, **kw):
        parent = self.styles.get(kw.pop('parent','Normal'), self.styles['Normal'])
        self.styles.add(ParagraphStyle(name=name, parent=parent, **kw))

    def _create(self):
        C = ReportConfig
        self._add('ReportTitle', fontSize=C.FONT_TITLE, textColor=C.C_PRIMARY,
                  spaceAfter=C.SP_SM, alignment=TA_CENTER, fontName='Helvetica-Bold', leading=24)
        self._add('SectionHeader', fontSize=C.FONT_H1, textColor=C.C_WHITE,
                  spaceAfter=C.SP_XS, spaceBefore=C.SP_MD, fontName='Helvetica-Bold',
                  leading=16, backColor=C.C_SECONDARY, borderPadding=(4,6,4,6))
        self._add('SubsectionHeader', fontSize=C.FONT_H2, textColor=C.C_SECONDARY,
                  spaceAfter=C.SP_XS, spaceBefore=C.SP_SM, fontName='Helvetica-Bold', leading=13)
        self._add('ReportNormal', fontSize=C.FONT_NORMAL, textColor=C.C_DARK,
                  spaceAfter=C.SP_XS, leading=12)
        self._add('ExecutiveSummary', fontSize=C.FONT_NORMAL, textColor=C.C_DARK,
                  spaceAfter=C.SP_XS, leading=12)
        self._add('CriticalBanner', fontSize=12, textColor=C.C_WHITE,
                  spaceAfter=0, fontName='Helvetica-Bold', backColor=C.C_CRITICAL,
                  leading=15, borderPadding=(5,8,5,8))
        self._add('HighBanner', fontSize=12, textColor=C.C_WHITE,
                  spaceAfter=0, fontName='Helvetica-Bold', backColor=C.C_HIGH,
                  leading=15, borderPadding=(5,8,5,8))
        self._add('VulnCardLabel', fontSize=C.FONT_SMALL, textColor=C.C_GRAY,
                  fontName='Helvetica-Bold', leading=11, spaceAfter=1)
        self._add('VulnCardValue', fontSize=C.FONT_NORMAL, textColor=C.C_DARK,
                  leading=12, spaceAfter=C.SP_XS)
        self._add('VulnCardDesc', fontSize=C.FONT_NORMAL, textColor=C.C_DARK,
                  leading=12, spaceAfter=C.SP_XS, leftIndent=4)
        self._add('VulnCardRemedy', fontSize=C.FONT_NORMAL,
                  textColor=HexColor('#1a5276'), leading=12, spaceAfter=C.SP_XS,
                  leftIndent=4, fontName='Helvetica-Bold')
        self._add('ReportFooter', fontSize=C.FONT_TINY, textColor=C.C_GRAY,
                  alignment=TA_CENTER, spaceBefore=C.SP_SM)
        self._add('KeyFinding', fontSize=C.FONT_NORMAL, textColor=C.C_DARK,
                  spaceAfter=2, leftIndent=8, leading=12)
        self._add('ListItem', fontSize=C.FONT_NORMAL, textColor=C.C_DARK,
                  spaceAfter=2, leftIndent=10, leading=12)
        self._add('Classification', fontSize=C.FONT_SMALL, textColor=C.C_CRITICAL,
                  alignment=TA_CENTER, fontName='Helvetica-Bold', spaceBefore=C.SP_SM)
        self._add('Disclaimer', fontSize=C.FONT_TINY, textColor=C.C_GRAY,
                  alignment=TA_CENTER, spaceBefore=2)

    def get(self, name):
        if not REPORTLAB_AVAILABLE: return None
        return self.styles.get(name, self.styles.get('Normal'))


# =============================================================================
# SECTION 5: TABLE BUILDER
# All tables use USABLE_WIDTH = 540pt = 7.5" so nothing overflows the margin.
# =============================================================================

class TableBuilder:
    W = ReportConfig.USABLE_WIDTH   # 540 pt

    def _base_cmds(self, hdr_color, font_size, hdr_size, C):
        return [
            ('GRID',         (0,0),(-1,-1), 0.4, C.C_GRAY_LT),
            ('BACKGROUND',   (0,0),(-1, 0), hdr_color),
            ('TEXTCOLOR',    (0,0),(-1, 0), C.C_WHITE),
            ('FONTNAME',     (0,0),(-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE',     (0,0),(-1, 0), hdr_size),
            ('TOPPADDING',   (0,0),(-1,-1), C.TBL_PAD),
            ('BOTTOMPADDING',(0,0),(-1,-1), C.TBL_PAD),
            ('LEFTPADDING',  (0,0),(-1,-1), 5),
            ('RIGHTPADDING', (0,0),(-1,-1), 5),
            ('FONTNAME',     (0,1),(-1,-1), 'Helvetica'),
            ('FONTSIZE',     (0,1),(-1,-1), font_size),
            ('ROWBACKGROUNDS',(0,1),(-1,-1),[C.C_WHITE, C.C_ROW_ALT]),
            ('VALIGN',       (0,0),(-1,-1), 'TOP'),
            ('ALIGN',        (0,0),(-1, 0), 'CENTER'),
            ('WORDWRAP',     (0,0),(-1,-1), 'LTR'),
        ]

    def make(self, data, col_widths=None, hdr_color=None, font_size=None,
             hdr_size=None, valign='TOP', extra_cmds=None):
        if not REPORTLAB_AVAILABLE: return None
        C = ReportConfig
        if hdr_color is None: hdr_color = C.C_PRIMARY
        if font_size  is None: font_size  = C.TBL_CELL_SZ
        if hdr_size   is None: hdr_size   = C.TBL_HDR_SZ
        if not data:
            return Table([['No data available']], colWidths=[self.W])
        tbl = Table(data, colWidths=col_widths, repeatRows=1)
        cmds = self._base_cmds(hdr_color, font_size, hdr_size, C)
        cmds.append(('VALIGN',(0,0),(-1,-1), valign))
        if extra_cmds: cmds += extra_cmds
        tbl.setStyle(TableStyle(cmds))
        return tbl

    # ---------- scan info [1.8" + 5.7" = 7.5"] ----------
    def scan_info_table(self, rd):
        if not REPORTLAB_AVAILABLE: return None
        C = ReportConfig
        risk = rd['summary']['overall_risk']
        risk_bg = {'Critical':C.C_CRITICAL,'High':C.C_HIGH,
                   'Medium':C.C_MEDIUM,'Low':C.C_LOW}.get(risk, C.C_GRAYINFO)
        scan = rd['scan_info']; sev = rd['summary']['vulnerability_by_severity']
        score = rd['summary']['security_score']
        data = [
            ['Target Domain',  TextUtils.safe_text(scan.get('target','N/A'), 400)],
            ['Scan ID',        str(scan.get('id','N/A'))],
            ['Scan Date',      str(scan.get('started_at','N/A'))[:19]],
            ['Status',         str(scan.get('status','N/A')).upper()],
            ['Overall Risk',   risk.upper()],
            ['Security Score', f"{score}/100"],
            ['Findings',       (f"Critical: {sev['critical']}  |  High: {sev['high']}  |  "
                                f"Medium: {sev['medium']}  |  Low: {sev['low']}  |  "
                                f"Info: {sev['info']}")],
            ['Report Date',    datetime.now().strftime('%Y-%m-%d %H:%M UTC')],
        ]
        W = self.W
        tbl = Table(data, colWidths=[1.8*inch, W-1.8*inch])
        cmds = [
            ('FONTNAME',      (0,0),(0,-1), 'Helvetica-Bold'),
            ('FONTNAME',      (1,0),(1,-1), 'Helvetica'),
            ('FONTSIZE',      (0,0),(-1,-1), 9),
            ('TOPPADDING',    (0,0),(-1,-1), 4),
            ('BOTTOMPADDING', (0,0),(-1,-1), 4),
            ('LEFTPADDING',   (0,0),(-1,-1), 6),
            ('RIGHTPADDING',  (0,0),(-1,-1), 6),
            ('GRID',          (0,0),(-1,-1), 0.5, C.C_GRAY_LT),
            ('BACKGROUND',    (0,0),(0,-1),  C.C_LIGHT),
            ('ROWBACKGROUNDS',(0,0),(-1,-1),[C.C_LIGHT, C.C_WHITE]*10),
            ('VALIGN',        (0,0),(-1,-1), 'MIDDLE'),
            ('WORDWRAP',      (1,0),(1,-1),  'LTR'),
            ('BACKGROUND',    (1,4),(1,4),   risk_bg),
            ('TEXTCOLOR',     (1,4),(1,4),   C.C_WHITE),
            ('FONTNAME',      (1,4),(1,4),   'Helvetica-Bold'),
        ]
        tbl.setStyle(TableStyle(cmds)); return tbl

    # ---------- KPI table [3.5" + 1.5" + 2.5" = 7.5"] ----------
    def kpi_table(self, data):
        if not REPORTLAB_AVAILABLE: return None
        C = ReportConfig; W = self.W
        cols = [W*0.467, W*0.200, W*0.333]
        status_colors = {
            'Action Required': (C.C_CRITICAL, C.C_WHITE),
            'Urgent':          (C.C_HIGH,     C.C_WHITE),
            'Review':          (C.C_WARNING,  C.C_WHITE),
            'Review Required': (C.C_WARNING,  C.C_WHITE),
            'None':            (C.C_SUCCESS,  C.C_WHITE),
            'Complete':        (C.C_SUCCESS,  C.C_WHITE),
            'Verified':        (C.C_SUCCESS,  C.C_WHITE),
        }
        extra = []
        for ri, row in enumerate(data[1:], 1):
            if len(row) >= 3:
                s = str(row[2]).strip()
                if s in status_colors:
                    bg, fg = status_colors[s]
                    extra += [('BACKGROUND',(2,ri),(2,ri),bg),
                               ('TEXTCOLOR', (2,ri),(2,ri),fg),
                               ('FONTNAME',  (2,ri),(2,ri),'Helvetica-Bold')]
        extra += [('ALIGN',(1,0),(2,-1),'CENTER')]
        return self.make(data, col_widths=cols, hdr_color=C.C_PRIMARY, extra_cmds=extra)

    # ---------- Risk table [1.5"+1.0"+1.2"+3.8" = 7.5"] ----------
    def risk_table(self, data, severity_counts):
        if not REPORTLAB_AVAILABLE: return None
        C = ReportConfig; W = self.W
        cols = [W*0.200, W*0.133, W*0.160, W*0.507]
        sc = {'CRITICAL':C.C_CRITICAL,'HIGH':C.C_HIGH,
              'MEDIUM':C.C_MEDIUM,'LOW':C.C_LOW,'INFO':C.C_GRAYINFO}
        extra = [('ALIGN',(1,0),(2,-1),'CENTER')]
        for ri, sev in enumerate(['CRITICAL','HIGH','MEDIUM','LOW','INFO'],1):
            if severity_counts.get(sev.lower(),0) > 0:
                bg = sc[sev]
                extra += [('BACKGROUND',(0,ri),(0,ri),bg),
                           ('TEXTCOLOR', (0,ri),(0,ri),C.C_WHITE),
                           ('FONTNAME',  (0,ri),(0,ri),'Helvetica-Bold')]
        return self.make(data, col_widths=cols, hdr_color=C.C_SECONDARY, extra_cmds=extra)

    # ---------- DNS table [1.5"+4.5"+1.5" = 7.5"] ----------
    def dns_table(self, dns_info):
        if not REPORTLAB_AVAILABLE: return None
        C = ReportConfig; W = self.W
        def ns(lst,n=3): return TextUtils.safe_text(', '.join(lst[:n]) if lst else 'Not Found',300)
        def st(val): return TextUtils.get_status_text('configured' if val else 'missing')
        data = [
            ['Record Type','Value','Status'],
            ['NS Records',   ns(dns_info.get('nameservers',[])),  st(dns_info.get('nameservers'))],
            ['MX Records',   ns(dns_info.get('mx_records',[])),   st(dns_info.get('mx_records'))],
            ['SPF Record',   TextUtils.safe_text(dns_info.get('spf') or 'Not Found',300),
                             st(dns_info.get('spf'))],
            ['DMARC Record', TextUtils.safe_text(dns_info.get('dmarc') or 'Not Found',300),
                             st(dns_info.get('dmarc'))],
            ['DNSSEC',       'Enabled' if dns_info.get('dnssec') else 'Not Enabled',
                             st(dns_info.get('dnssec'))],
            ['IPv6 (AAAA)',  ns(dns_info.get('ipv6',[]),2), st(dns_info.get('ipv6'))],
            ['CAA Record',   'Present' if dns_info.get('caa') else 'Not Found',
                             st(dns_info.get('caa'))],
        ]
        cols = [W*0.200, W*0.600, W*0.200]
        extra = [('ALIGN',(2,0),(2,-1),'CENTER')]
        for i in range(1,len(data)):
            s = data[i][2]
            if s == 'Missing':
                extra += [('BACKGROUND',(2,i),(2,i),C.C_DANGER),('TEXTCOLOR',(2,i),(2,i),C.C_WHITE),('FONTNAME',(2,i),(2,i),'Helvetica-Bold')]
            elif s in ('Configured','Enabled','Present'):
                extra += [('BACKGROUND',(2,i),(2,i),C.C_SUCCESS),('TEXTCOLOR',(2,i),(2,i),C.C_WHITE)]
            else:
                extra += [('BACKGROUND',(2,i),(2,i),C.C_WARNING),('TEXTCOLOR',(2,i),(2,i),C.C_WHITE)]
        return self.make(data, col_widths=cols, hdr_color=C.C_SUCCESS, extra_cmds=extra)

    # ---------- SSL table [2.0"+3.9"+1.6" = 7.5"] ----------
    def ssl_table(self, ssl_info):
        if not REPORTLAB_AVAILABLE: return None
        C = ReportConfig; W = self.W
        data = [
            ['SSL/TLS Property','Value','Assessment'],
            ['Certificate Issuer', ssl_info.get('issuer') or 'Unknown', 'Info'],
            ['TLS Versions',       ', '.join(ssl_info.get('tls_versions',[])) or 'Not Detected',
             'Secure' if any('1.2' in v or '1.3' in v for v in ssl_info.get('tls_versions',[])) else 'Review'],
            ['Deprecated TLS',    ', '.join(ssl_info.get('deprecated_tls',[])) or 'None',
             'VULNERABLE' if ssl_info.get('deprecated_tls') else 'OK'],
            ['Weak Ciphers',      ', '.join(ssl_info.get('weak_ciphers',[])[:4]) or 'None Detected',
             'FOUND' if ssl_info.get('weak_ciphers') else 'OK'],
            ['Wildcard Cert',     'Yes' if ssl_info.get('wildcard') else 'No',
             'Review' if ssl_info.get('wildcard') else 'OK'],
            ['Subject Alt Names', ', '.join(ssl_info.get('san',[])[:6]) or 'Not Found','Info'],
        ]
        cols = [W*0.267, W*0.520, W*0.213]
        ac = {'VULNERABLE':(C.C_CRITICAL,C.C_WHITE),'FOUND':(C.C_HIGH,C.C_WHITE),
              'Review':(C.C_WARNING,C.C_WHITE),'OK':(C.C_SUCCESS,C.C_WHITE),'Secure':(C.C_SUCCESS,C.C_WHITE)}
        extra = [('ALIGN',(2,0),(2,-1),'CENTER')]
        for i in range(1,len(data)):
            a = data[i][2]
            if a in ac:
                bg,fg=ac[a]
                extra += [('BACKGROUND',(2,i),(2,i),bg),('TEXTCOLOR',(2,i),(2,i),fg),('FONTNAME',(2,i),(2,i),'Helvetica-Bold')]
        return self.make(data, col_widths=cols, hdr_color=C.C_INFO, extra_cmds=extra)

    # ---------- Hosts table [2.2"+0.8"+2.0"+1.2"+1.3" = 7.5"] ----------
    def hosts_table(self, hosts):
        if not REPORTLAB_AVAILABLE: return None
        C = ReportConfig; W = self.W
        data = [['URL','Code','Page Title','IP Address','Technologies']]
        for h in hosts[:ReportConfig.MAX_HOSTS]:
            tech = ', '.join((h.get('tech') or [])[:3]) or 'N/A'
            code = str(h.get('status_code','N/A'))
            data.append([TextUtils.safe_text(h.get('url','N/A'),200),
                         code,
                         TextUtils.safe_text(h.get('title') or 'N/A',100),
                         h.get('ip','N/A'),
                         TextUtils.safe_text(tech,100)])
        cols = [W*0.293,W*0.107,W*0.267,W*0.160,W*0.173]
        extra = [('ALIGN',(1,0),(1,-1),'CENTER')]
        for i, row in enumerate(data[1:],1):
            cs = str(row[1])
            if cs.startswith('2'):
                extra += [('BACKGROUND',(1,i),(1,i),C.C_SUCCESS),('TEXTCOLOR',(1,i),(1,i),C.C_WHITE)]
            elif cs.startswith('3'):
                extra += [('BACKGROUND',(1,i),(1,i),C.C_WARNING),('TEXTCOLOR',(1,i),(1,i),C.C_WHITE)]
            elif cs.startswith(('4','5')):
                extra += [('BACKGROUND',(1,i),(1,i),C.C_DANGER),('TEXTCOLOR',(1,i),(1,i),C.C_WHITE)]
        return self.make(data, col_widths=cols, hdr_color=C.C_SUCCESS, extra_cmds=extra)

    # ---------- Ports table [1.3"+0.9"+0.6"+0.7"+1.0"+1.5"+0.7"+0.8" = 7.5"] ----------
    def ports_table(self, open_ports):
        if not REPORTLAB_AVAILABLE: return None
        C = ReportConfig; W = self.W
        data = [['Host','IP','Port','Proto','Service','Version','State','Risk']]
        for p in open_ports[:ReportConfig.MAX_PORTS]:
            data.append([TextUtils.safe_text(p.get('host','N/A'),60),
                         p.get('ip','N/A'), str(p.get('port','N/A')),
                         p.get('protocol','tcp'), p.get('service','unknown'),
                         TextUtils.safe_text(p.get('version') or 'N/A',80),
                         p.get('state','open'), p.get('risk','Low')])
        cols = [W*0.173,W*0.120,W*0.080,W*0.093,W*0.133,W*0.200,W*0.093,W*0.107]
        rc = {'High':(C.C_CRITICAL,C.C_WHITE),'Medium':(C.C_WARNING,C.C_WHITE),'Low':(C.C_SUCCESS,C.C_WHITE)}
        extra = []
        for i,row in enumerate(data[1:],1):
            r=str(row[7])
            if r in rc:
                bg,fg=rc[r]
                extra += [('BACKGROUND',(7,i),(7,i),bg),('TEXTCOLOR',(7,i),(7,i),fg),('FONTNAME',(7,i),(7,i),'Helvetica-Bold')]
        return self.make(data, col_widths=cols, hdr_color=C.C_WARNING,
                         font_size=7, extra_cmds=extra)


# =============================================================================
# SECTION 6: DATA PROCESSORS
# =============================================================================

class VulnerabilityProcessor:
    def count_by_severity(self, vulns):
        counts = {s:0 for s in ['critical','high','medium','low','info']}
        for v in vulns:
            sev = v.get('info',{}).get('severity','info').lower()
            counts[sev] = counts.get(sev,0)+1
        return counts

    def by_severity(self, vulns, sev):
        return [v for v in vulns if v.get('info',{}).get('severity','').lower()==sev]

    def extract_critical(self, v): return self.by_severity(v,'critical')
    def extract_high(self, v):     return self.by_severity(v,'high')
    def extract_medium(self, v):   return self.by_severity(v,'medium')
    def extract_low(self, v):      return self.by_severity(v,'low')

    def get_details(self, vuln):
        info = vuln.get('info',{}); cls = info.get('classification',{})
        refs = info.get('reference',[]); tags = info.get('tags',[])
        if isinstance(refs,str): refs=[refs]
        if isinstance(tags,str): tags=[tags]
        return {
            'name':         info.get('name','Unknown Vulnerability'),
            'severity':     info.get('severity','info').upper(),
            'description':  info.get('description','No description available.'),
            'remediation':  info.get('remediation','Refer to vendor advisory.'),
            'host':         vuln.get('host','N/A'),
            'matched_at':   vuln.get('matched-at', vuln.get('host','N/A')),
            'template_id':  vuln.get('template-id','N/A'),
            'template_url': vuln.get('template-url',''),
            'cve_id':       cls.get('cve-id',''),
            'cvss_score':   cls.get('cvss-score', cls.get('cvss-metrics','')),
            'cvss_vector':  cls.get('cvss-metrics',''),
            'cwe':          cls.get('cwe-id',''),
            'references':   refs[:5],
            'tags':         tags[:8],
            'matcher_name': vuln.get('matcher-name',''),
            'extracted':    vuln.get('extracted-results',[]),
        }

    def calculate_risk_score(self, vulns):
        if not vulns: return 100
        weights={'critical':25,'high':15,'medium':8,'low':2,'info':0}
        score=100; counts=self.count_by_severity(vulns)
        for sev,w in weights.items(): score-=counts.get(sev,0)*w
        return max(0,min(100,score))

    def get_risk_level(self, score):
        return ('Good' if score>=80 else 'Fair' if score>=60
                else 'Needs Improvement' if score>=40 else 'Critical')


class DNSSSLProcessor:
    def extract_dns_records(self, vulns):
        dns={'nameservers':[],'mx_records':[],'spf':None,'dmarc':None,
             'dnssec':False,'ipv6':[],'caa':False,'txt_records':[]}
        for v in vulns:
            try:
                tid=v.get('template-id',''); ex=v.get('extracted-results',[])
                if   'nameserver-fingerprint' in tid: dns['nameservers']=[n.rstrip('.') for n in ex]
                elif 'mx-fingerprint'         in tid: dns['mx_records']=ex
                elif 'spf-record-detect'      in tid: dns['spf']=ex[0] if ex else None
                elif 'dmarc-detect'           in tid: dns['dmarc']=ex[0] if ex else None
                elif 'dnssec-detection'       in tid: dns['dnssec']=True
                elif 'aaaa-fingerprint'       in tid: dns['ipv6']=ex
                elif 'caa-fingerprint'        in tid: dns['caa']=True
                elif 'txt-fingerprint'        in tid: dns['txt_records']=ex
            except: pass
        return dns

    def extract_ssl_info(self, vulns):
        ssl={'issuer':None,'san':[],'wildcard':False,
             'tls_versions':[],'weak_ciphers':[],'deprecated_tls':[]}
        for v in vulns:
            try:
                tid=v.get('template-id',''); ex=v.get('extracted-results',[])
                if   'ssl-issuer'         in tid: ssl['issuer']=ex[0] if ex else None
                elif 'ssl-dns-names'      in tid: ssl['san']=ex
                elif 'wildcard-tls'       in tid: ssl['wildcard']=True
                elif 'tls-version'        in tid: ssl['tls_versions'].extend(ex)
                elif 'weak-cipher-suites' in tid: ssl['weak_ciphers'].extend(ex)
                elif 'deprecated-tls'     in tid: ssl['deprecated_tls'].extend(ex)
            except: pass
        ssl['tls_versions']=list(set(ssl['tls_versions']))
        ssl['weak_ciphers']=list(set(ssl['weak_ciphers']))
        ssl['deprecated_tls']=list(set(ssl['deprecated_tls']))
        return ssl


class NmapProcessor:
    def extract_open_ports(self, nmap_results):
        ports=[]
        for n in nmap_results:
            try:
                ip=n.get('ip','N/A'); host=n.get('host','N/A')
                for p in n.get('open_ports',[]):
                    info={'ip':ip,'host':host,'port':p,'protocol':'tcp',
                          'service':'unknown','version':'','state':'open','risk':self._risk(p)}
                    for d in n.get('ports',[]):
                        if str(d.get('port',''))==str(p):
                            info.update({'protocol':d.get('protocol','tcp'),
                                         'service':d.get('service','unknown'),
                                         'version':d.get('version','')})
                            break
                    ports.append(info)
            except: pass
        return ports

    def _risk(self, p):
        if p in [21,22,23,25,3306,5432,1433,6379,27017,5000,8888,9200]: return 'High'
        if p in [80,443,8080,8443,3389,5900,8000]:                       return 'Medium'
        return 'Low'

    def extract_services(self, nmap_results):
        s=[]
        for n in nmap_results:
            for d in n.get('ports',[]):
                sv=d.get('service','')
                if sv and sv not in s: s.append(sv)
        return s


class TechnologyProcessor:
    def count_technologies(self, vulns):
        t=[]
        for v in vulns:
            if 'tech-detect' in v.get('template-id','') or 'waf-detect' in v.get('template-id',''):
                m=v.get('matcher-name','')
                if m and m not in t: t.append(m)
        return t[:15]

    def count_missing_headers(self, vulns):
        h=[]
        for v in vulns:
            if 'http-missing-security-headers' in v.get('template-id',''):
                hd=v.get('matcher-name','').replace('-',' ').title()
                if hd and hd not in h: h.append(hd)
        return h


class RiskCalculator:
    def generate_risk_matrix(self, report_data):
        vp=VulnerabilityProcessor()
        counts=vp.count_by_severity(report_data.get('vulnerabilities',[]))
        total=max(1,sum(counts.values()))
        return [
            ['Severity','Count','Percentage','Risk Description'],
            ['CRITICAL',str(counts['critical']),f"{counts['critical']/total*100:.1f}%",
             'Immediate remediation required -- system may be compromised'],
            ['HIGH',    str(counts['high']),    f"{counts['high']/total*100:.1f}%",
             'Urgent remediation -- high exploitation likelihood'],
            ['MEDIUM',  str(counts['medium']),  f"{counts['medium']/total*100:.1f}%",
             'Remediate within 30 days -- moderate exploitation risk'],
            ['LOW',     str(counts['low']),     f"{counts['low']/total*100:.1f}%",
             'Remediate at next maintenance window'],
            ['INFO',    str(counts['info']),    f"{counts['info']/total*100:.1f}%",
             'Informational -- no immediate action required'],
        ]


class RecommendationEngine:
    def __init__(self):
        self._t = {
            'headers': {
                'title':'Implement HTTP Security Headers',
                'description':'Missing security headers expose the application to XSS, clickjacking, MIME-sniffing, and data-theft attacks.',
                'effort':'Low','impact':'High','priority':'High',
                'actions':['Add Content-Security-Policy (CSP)','Add X-Frame-Options: DENY',
                           'Add X-Content-Type-Options: nosniff',
                           'Add Strict-Transport-Security (HSTS) with 1-year max-age',
                           'Add Referrer-Policy: strict-origin-when-cross-origin',
                           'Add Permissions-Policy to restrict browser features']},
            'tls': {
                'title':'Harden TLS Configuration',
                'description':'Deprecated TLS 1.0/1.1 and weak ciphers allow BEAST/POODLE downgrade attacks.',
                'effort':'Medium','impact':'High','priority':'Critical',
                'actions':['Disable TLS 1.0 and TLS 1.1 immediately',
                           'Enforce TLS 1.2 minimum; enable TLS 1.3',
                           'Remove NULL, EXPORT, RC4, DES, 3DES ciphers',
                           'Enable Forward Secrecy (ECDHE key exchange)',
                           'Validate certificate chain and enable OCSP stapling']},
            'dns': {
                'title':'Harden DNS / Email Configuration',
                'description':'Missing SPF, DMARC, or DNSSEC leaves the domain open to email spoofing and DNS hijacking.',
                'effort':'Low','impact':'Medium','priority':'High',
                'actions':['Publish SPF record: v=spf1 include:... -all',
                           'Publish DMARC with p=quarantine or p=reject',
                           'Enable DNSSEC on authoritative name servers',
                           'Add CAA records to restrict cert issuance']},
        }

    def generate_recommendations(self, rd):
        recs=[]
        if rd['summary'].get('missing_headers'):
            r=dict(self._t['headers']); r['specific']=rd['summary']['missing_headers']
            recs.append(r)
        ssl=rd['summary'].get('ssl_info',{})
        if ssl.get('deprecated_tls') or ssl.get('weak_ciphers'):
            r=dict(self._t['tls']); r['deprecated']=ssl.get('deprecated_tls',[])
            recs.append(r)
        dns=rd['summary'].get('dns_records',{})
        if not dns.get('spf') or not dns.get('dmarc'):
            recs.append(dict(self._t['dns']))
        return recs[:10]


# =============================================================================
# SECTION 11: MAIN REPORT GENERATOR
# =============================================================================

class ReportGenerator:
    def __init__(self):
        self.styles = ReportStyles()
        self.tbl    = TableBuilder()
        self.vp     = VulnerabilityProcessor()
        self.dssl   = DNSSSLProcessor()
        self.nmap   = NmapProcessor()
        self.tech   = TechnologyProcessor()
        self.rc     = RiskCalculator()
        self.reco   = RecommendationEngine()
        self.cfg    = ReportConfig()
        self.logger = LogUtils.get_logger()

    # -------------------------------------------------------
    # Data gathering
    # -------------------------------------------------------
    async def generate_report_data(self, db, scan_id):
        try:
            self.logger.log_info(f"Generating report data for scan: {scan_id}")
            scan = await db.get_scan(scan_id)
            if not scan: return None
            all_results = await db.get_results(scan_id)
            rd = {'scan_info':scan,'subdomains':[],'live_hosts':[],'wayback_urls':[],
                  'vulnerabilities':[],'nmap_results':[],'katana_endpoints':[],'summary':{}}
            seen_h=set(); seen_v=set()
            for result in all_results:
                try:
                    phase=result['phase']; data=result['data']
                    if isinstance(data,str):
                        try: data=json.loads(data)
                        except: continue
                    if phase=='subdomain':     rd['subdomains'].append(data)
                    elif phase=='httpx':
                        k=f"{data.get('url','')}-{data.get('ip','')}"
                        if k not in seen_h: seen_h.add(k); rd['live_hosts'].append(data)
                    elif phase=='wayback':     rd['wayback_urls'].append(data)
                    elif phase=='vulnerability':
                        k=(f"{data.get('template-id','')}-{data.get('host','')}"
                           f"-{data.get('matcher-name','')}")
                        if k not in seen_v: seen_v.add(k); rd['vulnerabilities'].append(data)
                    elif phase=='nmap':        rd['nmap_results'].append(data)
                    elif phase=='katana':      rd['katana_endpoints'].append(data)
                except Exception as e:
                    self.logger.log_error(f"Error processing result: {e}")
            rd['summary']=self._gen_summary(rd)
            return rd
        except Exception as e:
            self.logger.log_exception(f"Error generating report: {e}"); return None

    def _gen_summary(self, rd):
        vulns=rd['vulnerabilities']
        score=self.vp.calculate_risk_score(vulns)
        sev=self.vp.count_by_severity(vulns)
        risk=('Critical' if sev['critical']>0 else 'High' if sev['high']>0
              else 'Medium' if sev['medium']>0 else 'Low' if sev['low']>0 else 'Info')
        return {
            'total_subdomains':      len(rd['subdomains']),
            'total_live_hosts':      len(rd['live_hosts']),
            'total_vulnerabilities': len(vulns),
            'total_nmap_scans':      len(rd['nmap_results']),
            'total_endpoints':       len(rd['katana_endpoints']),
            'vulnerability_by_severity': sev,
            'missing_headers':       self.tech.count_missing_headers(vulns),
            'technologies_detected': self.tech.count_technologies(vulns),
            'dns_records':           self.dssl.extract_dns_records(vulns),
            'ssl_info':              self.dssl.extract_ssl_info(vulns),
            'open_ports':            self.nmap.extract_open_ports(rd['nmap_results']),
            'services_detected':     self.nmap.extract_services(rd['nmap_results']),
            'critical_vulns':        self.vp.extract_critical(vulns),
            'high_vulns':            self.vp.extract_high(vulns),
            'medium_vulns':          self.vp.extract_medium(vulns),
            'low_vulns':             self.vp.extract_low(vulns),
            'has_critical':          sev['critical']>0,
            'has_high':              sev['high']>0,
            'has_medium':            sev['medium']>0,
            'overall_risk':          risk,
            'security_score':        score,
            'risk_level':            self.vp.get_risk_level(score),
        }

    # -------------------------------------------------------
    # PDF generation
    # -------------------------------------------------------
    def generate_pdf(self, report_data, options=None):
        try:
            self.logger.log_info("Generating PDF report...")
            if not REPORTLAB_AVAILABLE:
                raise ImportError("ReportLab not available. pip install reportlab")
            if options is None: options=ReportConfig.DEFAULT_EXPORT_OPTIONS
            buf=io.BytesIO()
            doc=SimpleDocTemplate(buf, pagesize=self.cfg.PAGE_SIZE,
                                  rightMargin=self.cfg.MARGIN_RIGHT,
                                  leftMargin=self.cfg.MARGIN_LEFT,
                                  topMargin=self.cfg.MARGIN_TOP,
                                  bottomMargin=self.cfg.MARGIN_BOTTOM,
                                  title=f"Security Report - {report_data['scan_info']['target']}")
            elems=[]
            elems += self._cover(report_data)
            elems += self._exec_summary(report_data)
            elems.append(PageBreak())
            if report_data['summary']['has_critical'] or report_data['summary']['has_high']:
                elems += self._critical_high_section(report_data)
                elems.append(PageBreak())
            elems += self._all_vulns_table(report_data)
            elems.append(PageBreak())
            elems += self._assets_section(report_data)
            elems.append(PageBreak())
            elems += self._nmap_section(report_data)
            elems.append(PageBreak())
            elems += self._dns_ssl_section(report_data)
            if options.get('include_recommendations',True):
                elems.append(PageBreak())
                elems += self._recommendations_section(report_data)
            elems += self._footer(report_data)
            doc.build(elems)
            buf.seek(0)
            self.logger.log_info("PDF generated successfully")
            return buf.getvalue()
        except Exception as e:
            self.logger.log_exception(f"PDF generation failed: {e}"); raise

    # -------------------------------------------------------
    # Cover page
    # -------------------------------------------------------
    def _cover(self, rd):
        C=self.cfg; S=self.styles; elems=[]
        if rd['summary']['has_critical']:
            elems.append(Paragraph(
                "!  CRITICAL SECURITY RISK DETECTED  --  IMMEDIATE ACTION REQUIRED  !",
                S.get('CriticalBanner')))
        elif rd['summary']['has_high']:
            elems.append(Paragraph(
                "!  HIGH SECURITY RISK DETECTED  --  URGENT ATTENTION REQUIRED  !",
                S.get('HighBanner')))
        elems.append(Spacer(1,C.SP_SM))
        elems.append(Paragraph("SECURITY RECONNAISSANCE REPORT", S.get('ReportTitle')))
        elems.append(Paragraph("Recon Dashboard v4.1  |  Confidential", S.get('ReportNormal')))
        elems.append(Spacer(1,C.SP_SM))
        t=self.tbl.scan_info_table(rd)
        if t: elems.append(t)
        elems.append(Spacer(1,C.SP_SM))
        elems.append(Paragraph("CONFIDENTIAL -- FOR INTERNAL USE ONLY", S.get('Classification')))
        return elems

    # -------------------------------------------------------
    # Executive summary
    # -------------------------------------------------------
    def _exec_summary(self, rd):
        C=self.cfg; S=self.styles; elems=[]
        elems.append(Paragraph("EXECUTIVE SUMMARY", S.get('SectionHeader')))
        elems.append(Spacer(1,C.SP_XS))
        target=rd['scan_info'].get('target','N/A')
        sev=rd['summary']['vulnerability_by_severity']
        score=rd['summary']['security_score']
        risk=rd['summary']['overall_risk']
        elems.append(Paragraph(
            f"This report presents findings from a security reconnaissance of <b>{target}</b>. "
            f"The assessment covered subdomain enumeration, live host discovery, vulnerability "
            f"scanning (Nuclei), DNS/SSL analysis, and port scanning (Nmap). "
            f"Overall security score: <b>{score}/100</b>  --  Risk level: <b>{risk}</b>.",
            S.get('ExecutiveSummary')))
        elems.append(Spacer(1,C.SP_SM))

        elems.append(Paragraph("Assessment Metrics", S.get('SubsectionHeader')))
        kpi=[
            ['Metric','Count','Status'],
            ['Subdomains Discovered',   str(rd['summary']['total_subdomains']),      'Complete'],
            ['Live Hosts Identified',   str(rd['summary']['total_live_hosts']),       'Verified'],
            ['Total Vulnerabilities',   str(rd['summary']['total_vulnerabilities']),  'Review Required'],
            ['CRITICAL Issues',         str(sev['critical']),
             'None' if sev['critical']==0 else 'Action Required'],
            ['HIGH Issues',             str(sev['high']),
             'None' if sev['high']==0     else 'Urgent'],
            ['MEDIUM Issues',           str(sev['medium']),
             'None' if sev['medium']==0   else 'Review'],
            ['LOW Issues',              str(sev['low']),   'Minor'],
            ['Informational',           str(sev['info']),  'Awareness'],
            ['Open Ports Found',        str(len(rd['summary']['open_ports'])), 'Review'],
            ['Technologies Detected',   str(len(rd['summary']['technologies_detected'])),'Awareness'],
            ['Missing Sec Headers',     str(len(rd['summary']['missing_headers'])),
             'None' if not rd['summary']['missing_headers'] else 'Action Required'],
        ]
        k=self.tbl.kpi_table(kpi)
        if k: elems.append(k); elems.append(Spacer(1,C.SP_SM))

        elems.append(Paragraph("Risk Distribution", S.get('SubsectionHeader')))
        matrix=self.rc.generate_risk_matrix(rd)
        rt=self.tbl.risk_table(matrix,sev)
        if rt: elems.append(rt); elems.append(Spacer(1,C.SP_SM))

        elems.append(Paragraph("Key Findings", S.get('SubsectionHeader')))
        for i,f in enumerate(self._key_findings(rd)[:8],1):
            elems.append(Paragraph(f"{i}. {f}", S.get('KeyFinding')))
        return elems

    def _key_findings(self, rd):
        findings=[]
        for v in rd['summary']['critical_vulns'][:3]:
            d=self.vp.get_details(v)
            cve=f" ({d['cve_id']})" if d['cve_id'] else ''
            findings.append(f"[CRITICAL] {d['name']}{cve} -- {d['host']}")
        for v in rd['summary']['high_vulns'][:3]:
            d=self.vp.get_details(v)
            cve=f" ({d['cve_id']})" if d['cve_id'] else ''
            findings.append(f"[HIGH] {d['name']}{cve} -- {d['host']}")
        for v in rd['summary']['medium_vulns'][:2]:
            d=self.vp.get_details(v)
            findings.append(f"[MEDIUM] {d['name']} -- {d['host']}")
        if rd['summary']['missing_headers']:
            findings.append(f"Missing HTTP security headers: "
                            f"{', '.join(rd['summary']['missing_headers'][:4])}")
        if not findings:
            findings.append("No critical or high severity issues detected -- good security posture")
        return findings

    # -------------------------------------------------------
    # Critical / High detailed cards
    # -------------------------------------------------------
    def _critical_high_section(self, rd):
        C=self.cfg; S=self.styles; elems=[]
        crit=rd['summary']['critical_vulns']
        high=rd['summary']['high_vulns']

        if crit:
            elems.append(Paragraph("CRITICAL SEVERITY VULNERABILITIES", S.get('SectionHeader')))
            elems.append(Spacer(1,C.SP_XS))
            elems.append(Paragraph(
                f"<b>{len(crit)} critical severity issue(s)</b> identified. "
                f"These represent the highest risk and require <b>immediate remediation</b>. "
                f"Exploitation may result in full system compromise, data exfiltration, "
                f"or service disruption.",
                S.get('ExecutiveSummary')))
            elems.append(Spacer(1,C.SP_SM))
            for i,v in enumerate(crit[:ReportConfig.MAX_CH_CARDS],1):
                elems += self._vuln_card(v,i,'CRITICAL')
                elems.append(Spacer(1,C.SP_SM))

        if high:
            elems.append(Paragraph("HIGH SEVERITY VULNERABILITIES", S.get('SectionHeader')))
            elems.append(Spacer(1,C.SP_XS))
            elems.append(Paragraph(
                f"<b>{len(high)} high severity issue(s)</b> identified. "
                f"These carry significant exploitation risk and should be remediated "
                f"within 24-72 hours where possible.",
                S.get('ExecutiveSummary')))
            elems.append(Spacer(1,C.SP_SM))
            for i,v in enumerate(high[:ReportConfig.MAX_CH_CARDS],1):
                elems += self._vuln_card(v,i,'HIGH')
                elems.append(Spacer(1,C.SP_SM))
        return elems

    def _vuln_card(self, vuln, idx, severity):
        """Full-width styled card showing all available details for one vuln."""
        C=self.cfg; S=self.styles
        d=self.vp.get_details(vuln)
        W=self.tbl.W
        is_crit=(severity=='CRITICAL')
        hdr_bg  = C.C_CRITICAL if is_crit else C.C_HIGH
        card_bg = C.C_CRIT_BG  if is_crit else C.C_HIGH_BG
        border  = C.C_CRITICAL if is_crit else C.C_HIGH

        lbl_style   = S.get('VulnCardLabel')
        val_style   = S.get('VulnCardValue')
        desc_style  = S.get('VulnCardDesc')
        remedy_style= S.get('VulnCardRemedy')

        LW = 1.4*inch       # label column
        VW = W - LW         # value column (fills remaining width)

        rows=[]
        # Target host
        rows.append([Paragraph('<b>Target Host</b>', lbl_style),
                     Paragraph(TextUtils.safe_text(d['host'],400), val_style)])
        # Matched URL (if different)
        if d['matched_at'] and d['matched_at'] != d['host']:
            rows.append([Paragraph('<b>Matched URL</b>', lbl_style),
                         Paragraph(TextUtils.safe_text(d['matched_at'],400), val_style)])
        # CVE / CVSS / CWE
        if d['cve_id']:
            rows.append([Paragraph('<b>CVE ID</b>', lbl_style),
                         Paragraph(f"<b>{d['cve_id']}</b>", val_style)])
        if d['cvss_score']:
            rows.append([Paragraph('<b>CVSS Score</b>', lbl_style),
                         Paragraph(str(d['cvss_score']), val_style)])
        if d['cvss_vector']:
            rows.append([Paragraph('<b>CVSS Vector</b>', lbl_style),
                         Paragraph(TextUtils.safe_text(d['cvss_vector'],300), val_style)])
        if d['cwe']:
            rows.append([Paragraph('<b>CWE</b>', lbl_style),
                         Paragraph(str(d['cwe']), val_style)])
        if d['template_id'] and d['template_id']!='N/A':
            rows.append([Paragraph('<b>Template ID</b>', lbl_style),
                         Paragraph(TextUtils.safe_text(d['template_id'],200), val_style)])
        if d['tags']:
            rows.append([Paragraph('<b>Tags</b>', lbl_style),
                         Paragraph(', '.join(d['tags']), val_style)])
        if d['matcher_name']:
            rows.append([Paragraph('<b>Matcher</b>', lbl_style),
                         Paragraph(TextUtils.safe_text(d['matcher_name'],200), val_style)])
        # Description
        rows.append([Paragraph('<b>Description</b>', lbl_style),
                     Paragraph(TextUtils.safe_text(d['description'],1000), desc_style)])
        # Remediation
        rows.append([Paragraph('<b>Remediation</b>', lbl_style),
                     Paragraph(TextUtils.safe_text(d['remediation'],1000), remedy_style)])
        # Extracted results
        if d['extracted']:
            ex_txt='; '.join(TextUtils.safe_text(str(e),150) for e in d['extracted'][:5])
            rows.append([Paragraph('<b>Extracted Data</b>', lbl_style),
                         Paragraph(ex_txt, val_style)])
        # References
        if d['references']:
            ref_txt='  '.join(TextUtils.safe_text(str(r),200) for r in d['references'])
            rows.append([Paragraph('<b>References</b>', lbl_style),
                         Paragraph(ref_txt, val_style)])

        # Inner detail table
        detail_tbl=Table(rows, colWidths=[LW, VW])
        detail_tbl.setStyle(TableStyle([
            ('FONTNAME',     (0,0),(-1,-1),'Helvetica'),
            ('FONTSIZE',     (0,0),(-1,-1),8),
            ('TOPPADDING',   (0,0),(-1,-1),3),
            ('BOTTOMPADDING',(0,0),(-1,-1),3),
            ('LEFTPADDING',  (0,0),(-1,-1),5),
            ('RIGHTPADDING', (0,0),(-1,-1),5),
            ('VALIGN',       (0,0),(-1,-1),'TOP'),
            ('WORDWRAP',     (0,0),(-1,-1),'LTR'),
            ('BACKGROUND',   (0,0),(0,-1), card_bg),
            ('LINEBELOW',    (0,0),(-1,-2),0.3,C.C_GRAY_LT),
        ]))

        # Header paragraph
        hdr_para=Paragraph(
            f"#{idx}  {severity}  --  {TextUtils.safe_text(d['name'],200)}",
            S.get('CriticalBanner' if is_crit else 'HighBanner'))

        # Outer card table (header + details)
        outer=Table([[hdr_para],[detail_tbl]], colWidths=[W])
        outer.setStyle(TableStyle([
            ('BOX',          (0,0),(-1,-1),1.2,border),
            ('TOPPADDING',   (0,0),(-1,-1),0),
            ('BOTTOMPADDING',(0,0),(-1,-1),0),
            ('LEFTPADDING',  (0,0),(-1,-1),0),
            ('RIGHTPADDING', (0,0),(-1,-1),0),
            ('VALIGN',       (0,0),(-1,-1),'TOP'),
        ]))
        return [KeepTogether([outer])]

    # -------------------------------------------------------
    # All vulnerabilities table
    # -------------------------------------------------------
    def _all_vulns_table(self, rd):
        C=self.cfg; S=self.styles; W=self.tbl.W; elems=[]
        elems.append(Paragraph("ALL VULNERABILITY FINDINGS", S.get('SectionHeader')))
        elems.append(Spacer(1,C.SP_XS))

        sev=rd['summary']['vulnerability_by_severity']
        matrix=self.rc.generate_risk_matrix(rd)
        rt=self.tbl.risk_table(matrix,sev)
        if rt: elems.append(rt); elems.append(Spacer(1,C.SP_SM))

        # Medium  [2.0"+1.5"+2.2"+1.8" = 7.5"]
        med_cols=[W*0.267, W*0.200, W*0.293, W*0.240]
        if rd['summary']['medium_vulns']:
            elems.append(Paragraph("Medium Severity Findings", S.get('SubsectionHeader')))
            data=[['Vulnerability','Affected Host','Description','Remediation']]
            for v in rd['summary']['medium_vulns'][:20]:
                d=self.vp.get_details(v)
                cve=f" ({d['cve_id']})" if d['cve_id'] else ''
                data.append([TextUtils.safe_text(d['name']+cve,120),
                             TextUtils.safe_text(d['host'],100),
                             TextUtils.safe_text(d['description'],250),
                             TextUtils.safe_text(d['remediation'],200)])
            t=self.tbl.make(data,col_widths=med_cols,hdr_color=C.C_MEDIUM,font_size=7)
            if t: elems.append(t); elems.append(Spacer(1,C.SP_SM))

        # Low
        if rd['summary']['low_vulns']:
            elems.append(Paragraph("Low Severity Findings", S.get('SubsectionHeader')))
            data=[['Vulnerability','Affected Host','Description','Remediation']]
            for v in rd['summary']['low_vulns'][:15]:
                d=self.vp.get_details(v)
                data.append([TextUtils.safe_text(d['name'],120),
                             TextUtils.safe_text(d['host'],100),
                             TextUtils.safe_text(d['description'],250),
                             TextUtils.safe_text(d['remediation'],200)])
            t=self.tbl.make(data,col_widths=med_cols,hdr_color=C.C_LOW,font_size=7)
            if t: elems.append(t); elems.append(Spacer(1,C.SP_SM))

        # Info grouped  [1.5"+3.4"+2.6" = 7.5"]
        info_vulns=self.vp.by_severity(rd['vulnerabilities'],'info')
        if info_vulns:
            elems.append(Paragraph("Informational Findings", S.get('SubsectionHeader')))
            cats=defaultdict(list)
            for v in info_vulns[:30]:
                tid=v.get('template-id','')
                cat=('WAF Detection'     if 'waf-detect'               in tid else
                     'Missing Headers'   if 'missing-security-headers' in tid else
                     'DNS / Email'       if any(x in tid for x in ['spf','dmarc','mx-fingerprint']) else
                     'Technology Stack'  if any(x in tid for x in ['tech-detect','wappalyzer']) else
                     'SSL / TLS'         if any(x in tid for x in ['ssl','tls']) else
                     'Other')
                cats[cat].append(v)
            info_data=[['Category','Finding','Host']]
            for cat,vs in cats.items():
                for v in vs[:5]:
                    info_data.append([cat,
                        TextUtils.safe_text(v.get('info',{}).get('name','N/A'),100),
                        TextUtils.safe_text(v.get('host','N/A'),120)])
            t=self.tbl.make(info_data,col_widths=[W*0.200,W*0.453,W*0.347],
                            hdr_color=C.C_GRAYINFO,font_size=7)
            if t: elems.append(t)
        return elems

    # -------------------------------------------------------
    # Assets
    # -------------------------------------------------------
    def _assets_section(self, rd):
        C=self.cfg; S=self.styles; elems=[]
        elems.append(Paragraph("ASSET OVERVIEW", S.get('SectionHeader')))
        elems.append(Spacer(1,C.SP_XS))
        elems.append(Paragraph("DNS Configuration", S.get('SubsectionHeader')))
        t=self.tbl.dns_table(rd['summary']['dns_records'])
        if t: elems.append(t); elems.append(Spacer(1,C.SP_SM))
        if rd['live_hosts']:
            elems.append(Paragraph(f"Live Hosts  ({len(rd['live_hosts'])} total)",
                                   S.get('SubsectionHeader')))
            ht=self.tbl.hosts_table(rd['live_hosts'])
            if ht: elems.append(ht)
        return elems

    # -------------------------------------------------------
    # Nmap
    # -------------------------------------------------------
    def _nmap_section(self, rd):
        C=self.cfg; S=self.styles; elems=[]
        elems.append(Paragraph("NMAP PORT SCAN RESULTS", S.get('SectionHeader')))
        elems.append(Spacer(1,C.SP_XS))
        if rd['nmap_results'] and rd['summary']['open_ports']:
            elems.append(Paragraph(
                f"{len(rd['summary']['open_ports'])} open port(s) discovered across "
                f"{len(rd['nmap_results'])} host(s).",
                S.get('ReportNormal')))
            pt=self.tbl.ports_table(rd['summary']['open_ports'])
            if pt: elems.append(pt); elems.append(Spacer(1,C.SP_XS))
            if rd['summary']['services_detected']:
                elems.append(Paragraph(
                    f"<b>Detected Services:</b> "
                    f"{', '.join(rd['summary']['services_detected'][:15])}",
                    S.get('ReportNormal')))
        else:
            elems.append(Paragraph(
                "Nmap scan not executed in this cycle. Recommended for next assessment.",
                S.get('ReportNormal')))
        return elems

    # -------------------------------------------------------
    # DNS / SSL
    # -------------------------------------------------------
    def _dns_ssl_section(self, rd):
        C=self.cfg; S=self.styles; W=self.tbl.W; elems=[]
        elems.append(Paragraph("DNS & SSL / TLS CONFIGURATION", S.get('SectionHeader')))
        elems.append(Spacer(1,C.SP_XS))
        elems.append(Paragraph("SSL / TLS Analysis", S.get('SubsectionHeader')))
        st=self.tbl.ssl_table(rd['summary']['ssl_info'])
        if st: elems.append(st); elems.append(Spacer(1,C.SP_SM))

        if rd['summary']['technologies_detected']:
            elems.append(Paragraph("Detected Technologies", S.get('SubsectionHeader')))
            tech_data=[['#','Technology / Component']]
            for i,t in enumerate(rd['summary']['technologies_detected'],1):
                tech_data.append([str(i),t])
            tt=self.tbl.make(tech_data,col_widths=[W*0.067,W*0.933],
                             hdr_color=C.C_ACCENT,font_size=8)
            if tt: elems.append(tt); elems.append(Spacer(1,C.SP_SM))

        if rd['summary']['missing_headers']:
            elems.append(Paragraph("Missing HTTP Security Headers", S.get('SubsectionHeader')))
            hdr_risk={'Content Security Policy':'High','Strict Transport Security':'High',
                      'X Frame Options':'Medium','X Content Type Options':'Medium',
                      'Referrer Policy':'Low','Permissions Policy':'Low'}
            hd=[['#','Missing Header','Risk Level']]
            for i,h in enumerate(rd['summary']['missing_headers'],1):
                hd.append([str(i),h,hdr_risk.get(h,'Medium')])
            ht=self.tbl.make(hd,col_widths=[W*0.067,W*0.720,W*0.213],
                             hdr_color=C.C_DANGER,font_size=8)
            if ht: elems.append(ht)
        return elems

    # -------------------------------------------------------
    # Recommendations
    # -------------------------------------------------------
    def _recommendations_section(self, rd):
        C=self.cfg; S=self.styles; W=self.tbl.W; elems=[]
        elems.append(Paragraph("SECURITY RECOMMENDATIONS", S.get('SectionHeader')))
        elems.append(Spacer(1,C.SP_XS))
        recs=self.reco.generate_recommendations(rd)
        for idx,rec in enumerate(recs,1):
            elems.append(Paragraph(
                f"{idx}. {rec.get('title','Recommendation')}  "
                f"[Priority: {rec.get('priority','High')} | Effort: {rec.get('effort','Medium')}]",
                S.get('SubsectionHeader')))
            elems.append(Paragraph(rec.get('description',''), S.get('ReportNormal')))
            for action in rec.get('actions',[]):
                elems.append(Paragraph(f"  - {action}", S.get('ListItem')))
            elems.append(Spacer(1,C.SP_SM))

        if recs:
            elems.append(Paragraph("Suggested Remediation Timeline", S.get('SubsectionHeader')))
            plan=[
                ['Timeframe','Action Items'],
                ['Immediate\n(0-24 hrs)',
                 'Patch/mitigate all CRITICAL vulnerabilities. Disable deprecated TLS. Block high-risk open ports.'],
                ['Short-term\n(1-7 days)',
                 'Remediate HIGH findings. Implement missing security headers. Review SSL certificate chain.'],
                ['Medium-term\n(1-4 weeks)',
                 'Address MEDIUM findings. Harden DNS (SPF/DMARC/DNSSEC). Review technology exposure. Re-scan.'],
                ['Long-term\n(1-3 months)',
                 'Fix LOW findings. Establish continuous scanning pipeline. Security awareness training. Pen test.'],
            ]
            pt=self.tbl.make(plan,col_widths=[W*0.160,W*0.840],hdr_color=C.C_PRIMARY,font_size=8)
            if pt: elems.append(pt)
        return elems

    # -------------------------------------------------------
    # Footer
    # -------------------------------------------------------
    def _footer(self, rd):
        C=self.cfg; S=self.styles
        elems=[PageBreak(), Spacer(1,C.SP_SM)]
        elems.append(Paragraph(
            f"<b>Generated by:</b> Recon Dashboard v4.1  |  "
            f"<b>Timestamp:</b> {datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')}  |  "
            f"<b>Classification:</b> CONFIDENTIAL  |  "
            f"<b>Distribution:</b> Internal Security Team Only",
            S.get('ReportFooter')))
        elems.append(Paragraph(
            "This report is for authorised security assessment purposes only. "
            "Do not distribute without explicit written permission from the security team.",
            S.get('Disclaimer')))
        return elems

    # -------------------------------------------------------
    # Non-PDF formats
    # -------------------------------------------------------
    def generate_json(self, rd):
        try:
            self.logger.log_info("Generating JSON report...")
            return json.dumps(rd, indent=2, default=str)
        except Exception as e:
            self.logger.log_exception(f"JSON failed: {e}"); raise

    def generate_excel(self, rd):
        try:
            self.logger.log_info("Generating Excel report...")
            if not OPENPYXL_AVAILABLE:
                raise ImportError("OpenPyXL not available. pip install openpyxl")
            wb=Workbook()
            hf=XLFont(bold=True,color="FFFFFF",size=10)
            hfill=PatternFill(start_color="1a252f",end_color="1a252f",fill_type="solid")
            ha=XLAlignment(horizontal="center",vertical="center",wrap_text=True)
            ba=XLAlignment(wrap_text=True,vertical="top")
            ws=wb.active; ws.title="Summary"
            sev=rd['summary']['vulnerability_by_severity']
            summary_rows=[
                ["Metric","Value"],
                ["Target",          rd['scan_info']['target']],
                ["Scan ID",         rd['scan_info']['id']],
                ["Status",          rd['scan_info']['status']],
                ["Overall Risk",    rd['summary']['overall_risk']],
                ["Security Score",  f"{rd['summary']['security_score']}/100"],
                ["Subdomains",      rd['summary']['total_subdomains']],
                ["Live Hosts",      rd['summary']['total_live_hosts']],
                ["Total Vulns",     rd['summary']['total_vulnerabilities']],
                ["Critical",        sev['critical']],
                ["High",            sev['high']],
                ["Medium",          sev['medium']],
                ["Low",             sev['low']],
                ["Informational",   sev['info']],
            ]
            for ri,row in enumerate(summary_rows,1):
                for ci,val in enumerate(row,1):
                    cell=ws.cell(ri,ci,val)
                    if ri==1: cell.font=hf; cell.fill=hfill; cell.alignment=ha
                    else:     cell.alignment=ba
            ws.column_dimensions['A'].width=22
            ws.column_dimensions['B'].width=40
            # Vulns sheet
            ws2=wb.create_sheet("Vulnerabilities")
            vh=['Severity','Name','Host','CVE','CVSS','Template','Description','Remediation']
            for ci,h in enumerate(vh,1):
                c=ws2.cell(1,ci,h); c.font=hf; c.fill=hfill; c.alignment=ha
            for ri,v in enumerate(rd.get('vulnerabilities',[]),2):
                d=self.vp.get_details(v)
                for ci,val in enumerate([d['severity'],d['name'],d['host'],
                    d['cve_id'],str(d['cvss_score'] or ''),d['template_id'],
                    d['description'][:500],d['remediation'][:300]],1):
                    ws2.cell(ri,ci,str(val) if val else '').alignment=ba
            for cw,cl in zip([10,30,25,14,8,25,50,40],['A','B','C','D','E','F','G','H']):
                ws2.column_dimensions[cl].width=cw
            buf=io.BytesIO(); wb.save(buf); buf.seek(0)
            self.logger.log_info("Excel generated successfully")
            return buf.getvalue()
        except Exception as e:
            self.logger.log_exception(f"Excel failed: {e}"); raise

    def generate_docx(self, rd):
        try:
            self.logger.log_info("Generating Word report...")
            if not DOCX_AVAILABLE:
                raise ImportError("Python-Docx not available. pip install python-docx")
            doc=Document()
            t=doc.add_heading('Security Reconnaissance Report',0)
            t.alignment=WD_ALIGN_PARAGRAPH.CENTER
            doc.add_heading('Scan Information',level=1)
            tbl=doc.add_table(rows=6,cols=2); tbl.style='Table Grid'
            scan=rd['scan_info']; sev=rd['summary']['vulnerability_by_severity']
            for i,(lbl,val) in enumerate([
                ('Target',rd['scan_info'].get('target','N/A')),
                ('Scan ID',rd['scan_info'].get('id','N/A')),
                ('Status',rd['scan_info'].get('status','N/A')),
                ('Overall Risk',rd['summary']['overall_risk']),
                ('Started',str(rd['scan_info'].get('started_at','N/A'))[:19]),
                ('Completed',str(rd['scan_info'].get('completed_at','N/A'))[:19]),
            ]):
                tbl.cell(i,0).text=lbl; tbl.cell(i,1).text=str(val)
                tbl.cell(i,0).paragraphs[0].runs[0].bold=True
            doc.add_paragraph()
            doc.add_heading('Vulnerability Summary',level=1)
            for lbl,val in [('Total Vulnerabilities',rd['summary']['total_vulnerabilities']),
                            ('Critical',sev['critical']),('High',sev['high']),
                            ('Medium',sev['medium']),('Low',sev['low']),
                            ('Informational',sev['info'])]:
                doc.add_paragraph(f"{lbl}: {val}")
            buf=io.BytesIO(); doc.save(buf); buf.seek(0)
            self.logger.log_info("Word document generated successfully")
            return buf.getvalue()
        except Exception as e:
            self.logger.log_exception(f"Word failed: {e}"); raise

    def generate_csv(self, rd):
        try:
            self.logger.log_info("Generating CSV report...")
            buf=io.StringIO(); w=csv.writer(buf)
            w.writerow(['Severity','Name','Host','CVE','CVSS','Template','Description','Remediation'])
            for v in rd.get('vulnerabilities',[]):
                d=self.vp.get_details(v)
                w.writerow([d['severity'],d['name'],d['host'],d['cve_id'],
                            str(d['cvss_score'] or ''),d['template_id'],
                            d['description'][:400],d['remediation'][:300]])
            return buf.getvalue()
        except Exception as e:
            self.logger.log_exception(f"CSV failed: {e}"); raise

    def generate_xml(self, rd):
        try:
            self.logger.log_info("Generating XML report...")
            root=ET.Element('SecurityReport')
            si=ET.SubElement(root,'ScanInfo')
            ET.SubElement(si,'Target').text=rd['scan_info']['target']
            ET.SubElement(si,'ScanID').text=rd['scan_info']['id']
            ET.SubElement(si,'Status').text=rd['scan_info']['status']
            sm=ET.SubElement(root,'Summary')
            sev=rd['summary']['vulnerability_by_severity']
            for k in ['critical','high','medium','low','info']:
                ET.SubElement(sm,k.capitalize()).text=str(sev[k])
            vs=ET.SubElement(root,'Vulnerabilities')
            for v in rd.get('vulnerabilities',[])[:200]:
                d=self.vp.get_details(v)
                ve=ET.SubElement(vs,'Vulnerability')
                for tag,val in [('Severity',d['severity']),('Name',d['name']),
                                 ('Host',d['host']),('CVE',d['cve_id'] or ''),
                                 ('CVSS',str(d['cvss_score'] or '')),('Template',d['template_id'])]:
                    ET.SubElement(ve,tag).text=val
            return ET.tostring(root,encoding='unicode')
        except Exception as e:
            self.logger.log_exception(f"XML failed: {e}"); raise

    def generate_html(self, rd):
        try:
            self.logger.log_info("Generating HTML report...")
            sev=rd['summary']['vulnerability_by_severity']
            target=html_module.escape(rd['scan_info']['target'])
            risk=rd['summary']['overall_risk']
            score=rd['summary']['security_score']
            risk_color={'Critical':'#922b21','High':'#a04000',
                        'Medium':'#9a7d0a','Low':'#1a5276'}.get(risk,'#616a6b')
            rows=''
            for v in rd.get('vulnerabilities',[]):
                d=self.vp.get_details(v)
                sc={'CRITICAL':'#922b21','HIGH':'#a04000','MEDIUM':'#9a7d0a',
                    'LOW':'#1a5276','INFO':'#616a6b'}
                bg=sc.get(d['severity'],'#616a6b')
                rows+=(f"<tr><td style='background:{bg};color:#fff;font-weight:bold'>{d['severity']}</td>"
                       f"<td>{html_module.escape(d['name'])}</td>"
                       f"<td>{html_module.escape(d['host'])}</td>"
                       f"<td>{html_module.escape(d['cve_id'] or '')}</td>"
                       f"<td>{html_module.escape(str(d['cvss_score'] or ''))}</td></tr>")
            html_content=(
                f"<!DOCTYPE html><html lang='en'><head><meta charset='UTF-8'>"
                f"<title>Security Report - {target}</title><style>"
                f"body{{font-family:Arial,sans-serif;margin:30px;color:#1a252f}}"
                f"h1{{color:#1a252f;border-bottom:3px solid #2980b9;padding-bottom:6px}}"
                f"h2{{color:#2c3e50;margin-top:24px}}"
                f".badge{{display:inline-block;padding:4px 14px;border-radius:4px;"
                f"color:#fff;font-weight:bold;background:{risk_color}}}"
                f"table{{border-collapse:collapse;width:100%;margin-top:10px}}"
                f"th{{background:#1a252f;color:#fff;padding:7px 10px;text-align:left;font-size:13px}}"
                f"td{{border:1px solid #d5d8dc;padding:6px 10px;font-size:12px}}"
                f"tr:nth-child(even){{background:#f8f9fa}}</style></head><body>"
                f"<h1>Security Reconnaissance Report</h1>"
                f"<p><b>Target:</b> {target} &nbsp;|&nbsp;"
                f"<b>Risk:</b> <span class='badge'>{risk}</span> &nbsp;|&nbsp;"
                f"<b>Score:</b> {score}/100 &nbsp;|&nbsp;"
                f"<b>Generated:</b> {datetime.now().strftime('%Y-%m-%d %H:%M UTC')}</p>"
                f"<h2>Vulnerability Summary</h2>"
                f"<table><tr><th>Critical</th><th>High</th><th>Medium</th><th>Low</th><th>Info</th></tr>"
                f"<tr>"
                f"<td style='background:#922b21;color:#fff;font-weight:bold'>{sev['critical']}</td>"
                f"<td style='background:#a04000;color:#fff;font-weight:bold'>{sev['high']}</td>"
                f"<td style='background:#9a7d0a;color:#fff;font-weight:bold'>{sev['medium']}</td>"
                f"<td style='background:#1a5276;color:#fff'>{sev['low']}</td>"
                f"<td>{sev['info']}</td></tr></table>"
                f"<h2>All Findings ({rd['summary']['total_vulnerabilities']} total)</h2>"
                f"<table><tr><th>Severity</th><th>Name</th><th>Host</th><th>CVE</th><th>CVSS</th></tr>"
                f"{rows}</table></body></html>"
            )
            return html_content
        except Exception as e:
            self.logger.log_exception(f"HTML failed: {e}"); raise


# =============================================================================
# SINGLETON
# =============================================================================

report_generator = ReportGenerator()

if __name__ == "__main__":
    print("=" * 60)
    print("  RECON DASHBOARD REPORT GENERATOR  v4.1")
    print("=" * 60)
    print("  from services.report_generator import report_generator")
    print("  data = await report_generator.generate_report_data(db, scan_id)")
    print("  pdf  = report_generator.generate_pdf(data)")
    print("=" * 60)